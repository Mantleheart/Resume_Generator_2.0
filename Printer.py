from docx import Document
from docx.shared import Inches
from Demographic import getPair
from random import randrange, randint
from PySide6.QtWidgets import QProgressDialog, QWidget, QMessageBox
from PySide6 import QtCore
from docx.enum.text import WD_ALIGN_PARAGRAPH

import os
import csv
import datetime
import util

class loadingScreen(QWidget):
    def __init__(self):
        QWidget.__init__(self)


#add a new parameter for each pair that represents the root folder for that resume batch
def makeDoc(pair, mode, ID, parentFolder,params):

    #make a new sub-directory based on the ID
    #if mode is control, create the new child directory, or else reuse it

    headerColorationSeed = randrange(0,4)
    if mode == "test":
        headerColorationSeed = (headerColorationSeed + 1) % 4
    fonts = ['Ariel', 'Georgia', 'Helvetica','Avenir Next','Kefa','Times New Roman']
    
    #Decoration lists
    Emphasize = []
    Body = []
    Details = []
    
    #Reading Template

    localKeys = [
    "New York City High School",
    "Milwaukee High School",
    "Philadelphia High School",
    "Pittsburgh High School"
    ]
    
    localSchools = {
    "New York City High School":"Townsend Harris High School",
    "Milwaukee High School":"East High School",
    "Philadelphia High School":"Carver High School of Engineering & Science",
    "Pittsburgh High School":"Taylor Allderdice High School"
    }
    

  
    rangeUpper = util.countFiles("templates", "edu_template")
    docEdu = open('templates/edu_template' + str(randrange(1, rangeUpper)) + '.txt')
    docSchool = docEdu.read()
    docEdu.close()
    #replacing marker characters in edu_template
    docSchool = util.correctDates(docSchool.replace('#', pair[mode]['degree']).replace('&', pair[mode]['school']),params["BeginYearEdu"],params["EndYearEdu"],params["BeginMonthEdu"],params["EndMonthEdu"])
    

    
    #create document
    doc = Document()
    name = doc.add_heading('',headerColorationSeed)
    Emphasize.append(name.add_run(pair[mode]['fname'] + ' ' + pair[mode]['lname'].capitalize()))
    Info = doc.add_paragraph()
    
    #make random
    if mode == 'test':
        Info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    #if params["TestMode"] == "Before" and params["TestSection"] == "Address":
        #print("Inserting before address")
    #ADDRESS
    Details.append(Info.add_run(pair[mode]['address']))
    Details.append(Info.add_run('\r' + pair[mode]['number']))
    Details.append(Info.add_run('\r' + pair[mode]['email']))
    #END ADDRESS
    #if params["TestMode"] == "After" and params["TestSection"] == "Address":
        #print("Inserting before address")
    
    Emphasize.append(doc.add_heading('',headerColorationSeed).add_run('Education:'))
    
    
    #if params["TestMode"] == "Before" and params["TestSection"] == "Education":
        #print("Inserting before education")
    #EDUCATION
    Body.append(doc.add_paragraph().add_run(docSchool.replace('  ', ' ')))
    #END EDUCATION
    #if params["TestMode"] == "After" and params["TestSection"] == "Education":
        #print("Inserting after education")
    Emphasize.append(doc.add_heading('',headerColorationSeed).add_run('Professional Experience:'))

    if mode == 'test' and params["TestSection"] == "Work History":
        for each in localKeys:
            pair[mode]['work1'] = pair[mode]['work1'].replace(each, localSchools[each]);
    
    
    #if params["TestMode"] == "Before" and params["TestSection"] == "Work History":
        #print("Inserting before work")
    #WORK
    if pair[mode]['work1'] != '':
        work1 = doc.add_paragraph('')
        Body.append(work1.add_run(pair[mode]['work1']))
    if pair[mode]['work2'] != '':
        work2 = doc.add_paragraph('')
        Body.append(work2.add_run(pair[mode]['work2']))
    if pair[mode]['work3'] != '':
        work3 = doc.add_paragraph('')
        Body.append(work3.add_run(pair[mode]['work3']))
    #END WORK
    #if params["TestMode"] == "After" and params["TestSection"] == "Work History":
        #print("Inserting after work")
        
    Emphasize.append(doc.add_heading('',headerColorationSeed).add_run('Skills:'))
    
    #if params["TestMode"] == "Before" and params["TestSection"] == "Skills":
        #print("Inserting before skills")
    #SKILLS
    skills = doc.add_paragraph()
    Body.append(skills.add_run(pair[mode]['skills']))
    #END SKILLS
    #if params["TestMode"] == "After" and params["TestSection"] == "Skills":
        #print("Inserting after skills")

    fontSeed = randint(0,50)
    #change fonts and styles
    for each in Body:
        each.font.name = fonts[fontSeed % 6]
        
    for each in Details:
        each.font.name = fonts[fontSeed % 6]

    for each in Emphasize:
        each.font.name = fonts[fontSeed % 6]

    if mode == 'control':
        #make a new directory for the pair with the id as the name
        pairDirectoryPath = os.path.dirname(__file__) + "/" + parentFolder + "/" + str(ID)
        pairDirectory = os.mkdir(pairDirectoryPath)
    doc.save( parentFolder + '/' + str(ID) + '/' + pair[mode]['email'] + '.docx')
    return pair[mode]['email'], pair[mode]['gender'], pair[mode]['school'], str(datetime.datetime.now())

def generate(size, dir, path, outputName,params):
    currentTime = datetime.datetime.now()
    if len(outputName) <= 0:
        batchName = currentTime.strftime("%b-%d-%Y-%H%M%p%s")
    else:
        batchName = util.checkPath(dir,outputName)
    
    

    
    currentDirectory = os.getcwd()
    childDirectory = dir + "/" + batchName
    childDirectory = os.mkdir(childDirectory)
    record = [['email', 'gender', 'school', 'date']]
    ID = 1
    progressBarStep = int(size/10)
    progress = QProgressDialog("Generating Resumes","Close", 0, 100)
    progress.setWindowModality(QtCore.Qt.WindowModal)
    progress.setAutoClose(True)
    progress.show()



    while ID <= size:
        pair = getPair(path, params)
        record.append(makeDoc(pair, 'control', ID, batchName,params))
        record.append(makeDoc(pair, 'test', ID, batchName,params))
        ID += 1
        progress.setValue(ID)
    
    message = QMessageBox()
    message.setText("Resume generation complete.")
    message.exec()
    del pair
    return record
    
def begin(dir,path, outputName, params):

   
    with open('records.csv', 'w', newline='') as file:
        writer = csv.writer(file, delimiter=',', quotechar='"', quoting=csv.QUOTE_MINIMAL)
        writer.writerows(generate(100, dir, path, outputName,params))
#tkinter window initialized here.
#
#The main menu allows users to start generating. They add a name to the batch. 
#They can also choose the output directory
#
#They can quit or start over again when they're done.
#
#
